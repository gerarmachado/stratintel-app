import streamlit as st
import google.generativeai as genai
from openai import OpenAI
from docx import Document
from fpdf import FPDF
from io import BytesIO
import requests
from bs4 import BeautifulSoup
import os
import time
import datetime
from langchain_community.tools import DuckDuckGoSearchRun
import graphviz
import pypdf # AGREGE ESTA LIBRERÍA QUE FALTABA

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="StratIntel Solutions OS", page_icon="♟️", layout="wide")

# ==========================================
# 🔐 SISTEMA DE LOGIN
# ==========================================
def check_password():
    """Retorna True si el usuario/contraseña son correctos."""
    
    # Si no hay secretos configurados en la nube, permitimos acceso (Modo Desarrollo)
    if "passwords" not in st.secrets:
        st.warning("⚠️ Modo Desarrollo: No se detectó configuración de [passwords] en Secrets.")
        return True

    def password_entered():
        if st.session_state["username"] in st.secrets["passwords"] and \
           st.session_state["password"] == st.secrets["passwords"][st.session_state["username"]]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if st.session_state.get("password_correct", False):
        return True

    st.markdown("## ♟️ StratIntel Solutions: Acceso Restringido")
    st.text_input("Usuario", key="username")
    st.text_input("Contraseña", type="password", on_change=password_entered, key="password")
    
    if "password_correct" in st.session_state and not st.session_state["password_correct"]:
        st.error("❌ Credenciales inválidas")
    return False

if not check_password():
    st.stop()  

# ==========================================
# 🧠 BASE DE DATOS MAESTRA (GRAND UNIFIED STRATEGY)
# ==========================================
DB_CONOCIMIENTO = {
    "✨ RECOMENDACIÓN AUTOMÁTICA": {
        "desc": "La IA decide la mejor estrategia basándose en el contenido.",
        "preguntas": ["Identifica los hallazgos estratégicos más críticos.", "Realiza una evaluación integral de riesgos.", "Genera un Resumen Ejecutivo (BLUF).", "¿Cuáles son las anomalías o patrones ocultos más relevantes?"]
    },

    # =========================================================================
    # 🌍 BLOQUE 1: ESTRUCTURA, HISTORIA Y PODER (REFINADO)
    # =========================================================================
    
    # 1.1 META-TEORÍA Y FUERZAS PROFUNDAS -------------------------------------
    "--- 1.1 FUERZAS PROFUNDAS Y TEORÍA ---": { "desc": "", "preguntas": [] },

    "Renouvin & Duroselle (Fuerzas Profundas)": {
        "desc": "Las fuerzas subyacentes (geografía, demografía, economía) vs la decisión del estadista.",
        "preguntas": [
            "Fuerzas Profundas Materiales: ¿Cómo la geografía, demografía o economía limitan inevitablemente la acción política (Determinismo)?",
            "Fuerzas Profundas Espirituales: ¿Qué papel juegan los nacionalismos, el sentimiento colectivo o la psicología de masas en este conflicto?",
            "El Estadista vs la Fuerza: ¿El líder está moldeando la historia o simplemente está siendo arrastrado por corrientes profundas que no controla?"
        ]
    },
    "Dougherty & Pfaltzgraff (Teorías en Pugna)": {
        "desc": "Marco comparativo de teorías (Ambientales, Sistémicas, Decisorias).",
        "preguntas": [
            "Teorías Ambientales: ¿El conflicto es inevitable debido a la escasez de recursos o presiones ecológicas?",
            "Nivel de Análisis: ¿La causa raíz está en el Individuo (líder), el Estado (régimen) o el Sistema (anarquía)?",
            "Integración vs Desintegración: ¿Las fuerzas tecnológicas y económicas están uniendo a los actores o fragmentándolos en bloques hostiles?"
        ]
    },
    
    # 1.2 GRAN ESTRATEGIA Y CICLOS IMPERIALES ---------------------------------
    "Jean-Baptiste Duroselle (Todo Imperio Perecerá)": {
        "desc": "Ciclos vitales de las entidades políticas y sus causas de muerte.",
        "preguntas": [
            "Causa Exógena vs Endógena: ¿La amenaza principal proviene de una invasión externa o de la descomposición interna?",
            "Pérdida de Energía Creadora: ¿La sociedad ha dejado de innovar y se ha vuelto rígida y burocrática?",
            "Expansión Incontrolada: ¿Ha superado el Estado su 'radio de acción eficaz', volviéndose ingobernable?"
        ]
    },
    "Robert Kaplan (La Venganza de la Geografía)": {
        "desc": "El mapa como destino y las restricciones físicas del poder.",
        "preguntas": [
            "El Mapa del Alivio: ¿Cómo las montañas, ríos y desiertos imponen límites físicos insuperables a la ideología política?",
            "Zonas de Choque: ¿Está el conflicto ocurriendo en una 'zona de aplastamiento' (shatterbelt) histórica inevitable?",
            "Imperativos Geográficos: ¿Qué acciones está tomando el Estado simplemente porque su geografía se lo exige (salida al mar, defensa de fronteras llanas)?"
        ]
    },
    "Paul Kennedy (Auge y Caída de las Grandes Potencias)": {
        "desc": "Sobrestiramiento imperial (Imperial Overstretch).",
        "preguntas": [
            "Sobrestiramiento Imperial: ¿Están los compromisos militares y estratégicos del actor superando su capacidad económica para sostenerlos?",
            "Base Económica vs Poder Militar: ¿Se está sacrificando la inversión productiva interna para financiar la seguridad externa?",
            "Erosión Relativa: ¿El declive es absoluto o simplemente los rivales están creciendo más rápido?"
        ]
    },
    "E.H. Carr (La Crisis de los Veinte Años)": {
        "desc": "La crítica al utopismo liberal y la realidad del poder.",
        "preguntas": [
            "La Armonía de Intereses (Ilusión): ¿Están las potencias dominantes disfrazando sus propios intereses egoístas como 'valores universales' o 'bien común'?",
            "El Elemento Poder: ¿Se está ignorando que la política es, en última instancia, una lucha por el poder y no una discusión ética?",
            "Moralidad del Estado: ¿La moralidad que se predica es consistente con la capacidad real de ejercerla?"
        ]
    },
    "Ray Cline (La Fórmula del Poder)": {
        "desc": "Pp = (C + E + M) * (S + W) -> Potencial de Poder.",
        "preguntas": [
            "Masa Crítica (C): Evalúa Población + Territorio. ¿Tiene el estado la base física suficiente?",
            "Capacidad Económica y Militar (E + M): ¿Cuál es su fuerza bruta tangible?",
            "Estrategia y Voluntad (S + W): Estos son los multiplicadores. ¿Tiene el país una estrategia clara (S)? ¿Tiene el pueblo la voluntad nacional (W) de luchar? (Si esto es cero, el poder total es cero)."
        ]
    },
    
    # 1.3 REALISMO PURO Y GEOPOLÍTICA DURA ------------------------------------
    "Halford Mackinder (Teoría del Heartland)": {
        "desc": "El control de la Isla Mundial y el Pivote Geográfico.",
        "preguntas": [
            "Pivote Geográfico: ¿Quién controla actualmente el 'Heartland' (Eurasia central)?",
            "Cinturón Interior: ¿Hay un conflicto por el control de las zonas costeras que rodean el Heartland?",
            "Contención Terrestre: ¿Se está usando el poder terrestre para negar el acceso a las potencias marítimas?"
        ]
    },
    "Nicholas Spykman (Teoría del Rimland)": {
        "desc": "El control de los bordes anfibios (Rimland).",
        "preguntas": [
            "Anfibia Estratégica: Analiza el conflicto en las zonas costeras/peninsulares (Rimland).",
            "Cerco: ¿Están las potencias tratando de rodear al actor central desde el mar?",
            "Valor de las Alianzas: ¿Qué alianzas en el borde euroasiático son vitales para mantener el equilibrio?"
        ]
    },
    "Hans Morgenthau (Realismo Clásico Integral)": {
        "desc": "Los 6 Principios del Realismo Político y el Interés como Poder.",
        "preguntas": [
            "Leyes Objetivas: ¿Qué fuerzas inherentes a la naturaleza humana (egoísmo, dominio) están impulsando este conflicto?",
            "Interés y Poder: Define el 'Interés Nacional' de los actores en términos de poder, no de moralidad.",
            "Supervivencia del Estado: ¿Está la integridad territorial o política del Estado en riesgo directo?",
            "Autonomía de la Esfera Política: Analiza la decisión desde una lógica puramente política, ignorando consideraciones económicas o legales secundarias."
        ]
    },
    "Hans Morgenthau (El Otro Gran Debate: Interés Nacional)": {
        "desc": "Detección de la disolución del Interés Nacional por presiones Supranacionales (Moralismo Global) e Infranacionales (Intereses Sectoriales).",
        "preguntas": [
            "Desviación Supranacional (Utopismo): ¿Se está sacrificando la seguridad vital del Estado en nombre de 'principios morales universales', derecho internacional abstracto u organismos globales que no garantizan reciprocidad?",
            "Secuestro Infranacional (Captura del Estado): ¿La política exterior está siendo dictada por grupos subnacionales (lobbies corporativos, minorías étnicas, facciones ideológicas) que disfrazan su beneficio particular como 'Interés Nacional'?",
            "La Falacia Legalista: ¿Se está intentando resolver un problema político de poder mediante fórmulas legales o tratados que el adversario no respetará?",
            "Racionalidad vs Sentimentalismo: ¿La decisión es el resultado de un cálculo racional de poder o una respuesta emocional para satisfacer a la opinión pública doméstica?"
        ]
    },
    "Kenneth Waltz (Neorrealismo / Imágenes)": {
        "desc": "Las Tres Imágenes (Hombre, Estado, Sistema) y la Estructura Anárquica.",
        "preguntas": [
            "Tercera Imagen (Sistémica): ¿Cómo la anarquía internacional y la distribución de poder (polaridad) obligan al actor a actuar así?", 
            "Polaridad: ¿Cómo afecta la distribución de capacidades (unipolar/multipolar)?",
            "Segunda Imagen (Estatal): ¿Es el régimen político interno irrelevante para la política exterior en este caso?",
            "Equilibrio de Poder: ¿Está el actor haciendo 'Balancing' (aliarse contra el fuerte) o 'Bandwagoning' (unirse al fuerte)?",
            "Principio de Autoayuda: ¿Qué medidas unilaterales está tomando el actor para garantizar su propia seguridad? ¿El comportamiento es defensivo (seguridad) u ofensivo (poder)?"
        ]
    },
    "John Mearsheimer (Realismo Ofensivo)": {
        "desc": "La Tragedia de las Grandes Potencias y la Hegemonía.",
        "preguntas": [
            "Búsqueda de Hegemonía: ¿Está el actor intentando convertirse en el Hegemon regional para asegurar su supervivencia? ¿Está aprovechando oportunidades para alterar el status quo?",
            "Poder Detenedor del Agua: ¿Cómo la geografía (océanos, montañas) limita la proyección de poder del actor? Evalúa el potencial de poder latente (economía/población) vs poder militar actual.",
            "Maximizador de Poder: ¿Está el actor aprovechando cada oportunidad para debilitar a sus rivales potenciales? ¿Cómo está maximizando su poder relativo a expensas de sus vecinos?",
            "Estrategia de 'Buck-Passing': ¿Está intentando que otro estado asuma el costo de contener al agresor?"
        ]
    },
    "Stephen Walt & Robert Jervis (Realismo Defensivo)": {
        "desc": "Equilibrio de Amenazas y Dilema de Seguridad.",
        "preguntas": [
            "Teoría del Equilibrio de Amenazas: Evalúa la amenaza combinando: 1) Poder Agregado, 2) Geografía, 3) Capacidad Ofensiva, 4) Intenciones Agresivas. ¿Quién es percibido como el más amenazante (no solo el más fuerte)?",
            "Dilema de Seguridad: ¿Las medidas defensivas de un actor están siendo malinterpretadas como ofensivas por el otro?",
            "Espiral de Conflicto: ¿Cómo una acción defensiva ha provocado una reacción hostil involuntaria? ¿Las intenciones agresivas son reales o producto de la incertidumbre sistémica?"
        ]
    },
    "Realismo Neoclásico (Schweller)": {
        "desc": "El sistema presiona, pero la política interna decide.",
        "preguntas": [
            "¿Qué variables domésticas están filtrando o bloqueando la respuesta al sistema internacional?",
            "¿Es el estado 'coherente' o están las élites fragmentadas?",
            "¿Tiene el gobierno la capacidad extractiva para movilizar recursos ante la amenaza?"
        ]
    },
    "Realismo Periférico (Carlos Escudé)": {
        "desc": "Estrategia de supervivencia para estados dependientes (Sur Global).",
        "preguntas": [
            "Costo-Beneficio de la Soberanía: ¿El costo de confrontar al Hegemon supera los beneficios para el bienestar ciudadano?",
            "Política de Alineamiento: ¿Debería el estado adoptar un perfil bajo o alinearse para obtener recursos y evitar sanciones?",
            "Evaluación de Autonomía: ¿Se está sacrificando el desarrollo económico por una retórica nacionalista vacía?"
        ]
    },

    # -------------------------------------------------------------------------
    # 🤝 BLOQUE 2: ESCUELA LIBERAL Y CONSTRUCTIVISTA (INSTITUCIONES E IDENTIDAD)
    # -------------------------------------------------------------------------
    "--- LIBERALISMO, IDENTIDAD ---": { "desc": "", "preguntas": [] },

    "Joseph Nye (Poder Multidimensional 3D)": {
        "desc": "Soft Power, Smart Power y el Tablero de Ajedrez Tridimensional.",
        "preguntas": [
            "Dimensión Soft Power: ¿Qué activos de cultura, valores o políticas otorgan atracción y legitimidad al actor?",
            "Dimensión Smart Power: ¿Está combinando eficazmente la coerción (Hard) con la persuasión (Soft)?",
            "Tablero Superior (Militar): Analiza la distribución de poder militar (¿Unipolar?).",
            "Tablero Medio (Económico): Analiza la distribución económica (¿Multipolar?).",
            "Tablero Inferior (Transnacional): ¿Qué actores no estatales (Hackers, ONGs, Terrorismo) actúan fuera del control estatal?"
        ]
    },
    "Immanuel Kant (Triángulo de la Paz Liberal)": {
        "desc": "Paz Democrática, Interdependencia Económica e Instituciones.",
        "preguntas": [
            "Paz Democrática: ¿Son los actores democracias? (Si lo son, la probabilidad de guerra disminuye drásticamente).",
            "Interdependencia Económica: ¿El nivel de comercio mutuo hace que la guerra sea demasiado costosa?",
            "Organizaciones Internacionales: ¿Pertenecen a instituciones comunes que medien el conflicto?",
            "Derecho Cosmopolita: ¿Existe un respeto supranacional por los derechos de los ciudadanos?"
        ]
    },
    "Keohane & Nye (Neoliberalismo Institucional)": {
        "desc": "Interdependencia Compleja y Regímenes Internacionales.",
        "preguntas": [
            "Canales Múltiples: ¿Existen conexiones entre sociedades (no solo entre gobiernos)? ¿Qué instituciones facilitan la cooperación?",
            "Ausencia de Jerarquía: ¿Están los temas militares subordinados a temas económicos o ecológicos en esta crisis?",
            "Interdependencia Compleja: ¿Los vínculos económicos hacen la guerra irracional?",
            "Regímenes Internacionales: ¿Qué normas o reglas implícitas gobiernan las expectativas? ¿Existe un régimen internacional que regule este conflicto?"
        ]
    },
    "Alexander Wendt (Constructivismo Social)": {
        "desc": "La anarquía es lo que los estados hacen de ella.",
        "preguntas": [
            "Culturas de la Anarquía: ¿El sistema es Hobbesiano (Enemigos), Lockeano (Rivales) o Kantiano (Amigos)?",
            "Estructura Ideacional: ¿Cómo las identidades históricas y normas sociales definen el interés nacional?",
            "Ciclo de Refuerzo: ¿Cómo las interacciones pasadas han construido la percepción actual de 'amenaza'?",
            "Normas Internacionales: ¿Qué normas están constriñendo o habilitando la acción?"
        ]
    },
    "Teoría de la Integración Económica (Etapas y Modelos)": {
        "desc": "Niveles de fusión de mercados (Balassa y otros).",
        "preguntas": [
            "Nivel de Integración: ¿En qué fase están? 1) Zona de Libre Comercio (eliminar aranceles), 2) Unión Aduanera (arancel externo común), 3) Mercado Común (movilidad de factores), 4) Unión Económica (política fiscal/monetaria).",
            "Creación vs Desviación de Comercio: ¿El acuerdo genera riqueza real o simplemente desplaza a proveedores más eficientes externos?",
            "Spillover (Desbordamiento): ¿La integración económica está forzando inevitablemente la integración política?"
        ]
    },
    "Esther Barbé (Multilateralismo y Potencias Emergentes)": {
        "desc": "Adaptación del orden internacional y contestación normativa.",
        "preguntas": [
            "Contestación Normativa: ¿Las potencias emergentes están desafiando las reglas del juego o solo quieren un asiento en la mesa?",
            "Multilateralismo a la Carta: ¿Están los actores eligiendo selectivamente qué normas cumplir y cuáles ignorar?",
            "Cambio de Poder: ¿Las instituciones actuales reflejan la distribución real de poder o están obsoletas?"
        ]
    },
    "Samuel Huntington (Choque de Civilizaciones)": {
        "desc": "Conflictos de identidad cultural y religiosa.",
        "preguntas": [
            "Líneas de Falla: ¿Ocurre el conflicto en la frontera entre dos civilizaciones distintas?",
            "Núcleo Identitario: ¿Es el núcleo del conflicto la identidad religiosa o cultural?",
            "Síndrome del País Pariente (Kin-Country): ¿Están otros estados interviniendo por lealtad cultural/religiosa?",
            "Occidente vs El Resto: ¿Es una reacción contra la imposición de valores occidentales?"
        ]
    },

    # =========================================================================
    # ⚔️ BLOQUE 3: ESTRATEGIA MILITAR Y TRANSFORMACIÓN DE LA GUERRA
    # =========================================================================
    "--- ARTE DE LA GUERRA Y NUEVOS CONFLICTOS ---": { "desc": "", "preguntas": [] },

    "B.H. Liddell Hart (La Estrategia de Aproximación Indirecta)": {
        "desc": "Evitar la fortaleza, atacar la debilidad, dislocar al enemigo.",
        "preguntas": [
            "Línea de Menor Resistencia: ¿Está el actor atacando donde el enemigo menos lo espera (física o psicológicamente)?",
            "Dislocación: ¿Las maniobras han logrado separar al enemigo de su base, suministros o equilibrio mental antes del combate?",
            "Objetivos Alternativos: ¿Tiene el plan flexibilidad para cambiar de objetivo y mantener al enemigo en dilema?"
        ]
    },
    "Martin van Creveld (La Transformación de la Guerra)": {
        "desc": "Guerra No-Trinitaria y conflictos de baja intensidad.",
        "preguntas": [
            "Ruptura de la Trinidad: ¿El conflicto ignora la distinción clásica entre Gobierno, Ejército y Pueblo?",
            "Actores No Estatales: ¿Son las facciones, tribus o señores de la guerra más relevantes que el Estado?",
            "Guerra por la Existencia: ¿Se lucha por intereses políticos racionales o por mera supervivencia e identidad?"
        ]
    },
    "Mary Kaldor (Las Nuevas Guerras)": {
        "desc": "Conflictos post-Guerra Fría: Identidad + Globalización + Criminalidad.",
        "preguntas": [
            "Política de Identidad: ¿Se moviliza a la gente basándose en etiquetas étnicas/religiosas en lugar de ideología?",
            "Métodos de Terror: ¿Es el desplazamiento forzado y el ataque a civiles el objetivo central, no un daño colateral?",
            "Economía Depredadora: ¿Se financia la guerra mediante saqueo, mercado negro o ayuda humanitaria desviada?"
        ]
    },
    "Sun Tzu (El Arte de la Guerra)": {
        "desc": "Engaño, velocidad y victoria sin combate.",
        "preguntas": [
            "El Engaño: ¿Toda la operación se basa en una finta o distracción?",
            "Ganar sin luchar: ¿Está el actor logrando sus objetivos políticos sin uso cinético de fuerza?",
            "Conocimiento: ¿Conoce el actor al enemigo y a sí mismo?", 
            "Terreno: ¿Es el terreno mortal, disperso o clave? ¿Cómo afecta la maniobra?"
        ]
    },
    "Carl von Clausewitz (La Guerra Absoluta)": {
        "desc": "La guerra como continuación de la política.",
        "preguntas": [
            "Trinidad Paradójica: Analiza la relación entre Pasión (Pueblo), Probabilidad (Ejército) y Razón (Gobierno).",
            "Niebla y Fricción: ¿Qué imprevistos están ralentizando la operación?",
            "Centro de Gravedad (COG): ¿Cuál es la fuente de poder del enemigo que, si cae, todo el sistema colapsa?",
            "Política: ¿Es esta acción militar coherente con el objetivo político final?"
        ]
    },
    "Guerra Híbrida (Doctrina Gerasimov)": {
        "desc": "Sincronización de medios militares y no militares.",
        "preguntas": [
            "Fase Latente: ¿Se usa desinformación para desestabilizar antes del conflicto?",
            "Fuerzas Proxy: ¿Se utilizan actores no estatales para negar responsabilidad?",
            "Guerra Económica/Informativa: ¿Es el ataque principal cinético (bombas) o no cinético (sanciones/hackeos)?",
            "Dominio de la Información: ¿Es el ataque informativo más devastador que el físico?"
        ]
    },
    "Qiao Liang & Wang Xiangsui (Guerra Irrestricta)": {
        "desc": "Todo es un arma: leyes, economía, drogas, medios.",
        "preguntas": [
            "Desbordamiento del Campo de Batalla: ¿Se está usando el sistema legal (Lawfare) como arma?",
            "Guerra Financiera: ¿Se están atacando las monedas o mercados del adversario?",
            "Guerra Cultural: ¿Se están atacando los valores fundacionales de la sociedad objetivo?"
        ]
    },

    # =========================================================================
    # 💰 BLOQUE 4: GEOECONOMÍA, TRANSNACIONALISMO Y ANARQUÍA
    # =========================================================================
    "--- ECONOMÍA ILÍCITA Y CAOS ---": { "desc": "", "preguntas": [] },

    "Moisés Naím (Ilícito y el Fin del Poder)": {
        "desc": "El lado oscuro de la globalización y la erosión del Estado.",
        "preguntas": [
            "Las Cinco Guerras: Analiza el tráfico de: 1) Drogas, 2) Armas, 3) Personas, 4) Propiedad Intelectual, 5) Dinero sucio.",
            "Micropoderes: ¿Están actores pequeños y ágiles burlando las defensas de grandes burocracias estatales?",
            "Estado Hueco: ¿Tienen las instituciones la fachada de gobierno pero están carcomidas por redes criminales?"
        ]
    },
    "Robert Kaplan (La Anarquía que Viene)": {
        "desc": "Escasez, tribalismo y erosión de fronteras.",
        "preguntas": [
            "Estrés de Recursos: ¿Es la escasez de agua, tierra o comida el motor oculto del conflicto?",
            "Retribalización: ¿Están colapsando las identidades nacionales en favor de lealtades de clan o secta?",
            "Fronteras Porosas: ¿El mapa político oficial ha dejado de representar la realidad del control territorial?"
        ]
    },
    "Holm y Sorensen (Globalización Desigual)": {
        "desc": "¿De quién es el orden mundial? El fin de la Guerra Fría y la brecha Norte-Sur.",
        "preguntas": [
            "Ganadores y Perdedores: La globalización no es uniforme. ¿Quién se está integrando en el núcleo económico y quién está siendo marginado a la periferia irrelevante?",
            "Soberanía Fragmentada: ¿El Estado está perdiendo control frente a fuerzas globales (mercados) o frente a fuerzas locales (fragmentación étnica/regional)?",
            "El Dilema del Estado Débil: ¿Se está imponiendo un modelo de 'democracia liberal' en un estado que carece de las estructuras básicas para sostenerlo?"
        ]
    },
    "Edward Luttwak (Geoeconomía)": {
        "desc": "La lógica del conflicto con la gramática del comercio.",
        "preguntas": [
            "Armamentalización del Comercio: ¿Se usan aranceles o bloqueos como armas?",
            "Predación de Inversiones: ¿Está un estado adquiriendo infraestructura crítica del rival?",
            "Soberanía Tecnológica: ¿Se está bloqueando el acceso a tecnología clave?"
        ]
    },

    # =========================================================================
    # 🤝 BLOQUE 5: NEGOCIACIÓN, JUEGOS Y CONFLICTO
    # =========================================================================
    "--- ESTRATEGIA DE INTERACCIÓN ---": { "desc": "", "preguntas": [] },

    "Thomas Schelling (La Estrategia del Conflicto)": {
        "desc": "Disuasión, Compulsión y la Racionalidad de lo Irracional.",
        "preguntas": [
            "Compulsión vs Disuasión: ¿Se intenta impedir una acción (Disuasión) o forzar a que ocurra (Compulsión)?",
            "Puntos Focales (Schelling Points): ¿Existe una solución obvia donde convergerán las expectativas de ambos sin comunicarse?",
            "La Racionalidad de la Irracionalidad: ¿Se está fingiendo locura o descontrol para obligar al otro a ceder?",
            "Quemar los Barcos: ¿El actor se ha quitado a sí mismo la opción de retroceder para hacer creíble su amenaza?"
        ]
    },
    "William Ury (Cómo Negociar sin Ceder)": {
        "desc": "Negociación basada en principios y superación de bloqueos.",
        "preguntas": [
            "Intereses vs Posiciones: ¿Qué es lo que realmente quieren (Interés) vs lo que dicen que quieren (Posición)?",
            "MAPAN (BATNA): ¿Cuál es la Mejor Alternativa a un Acuerdo Negociado de cada parte? (Quién tiene más poder de retiro).",
            "Separar a la Persona del Problema: ¿Están las emociones o egos bloqueando la solución técnica?",
            "El Puente de Oro: ¿Se le está ofreciendo al adversario una salida digna para que no pierda la cara?"
        ]
    },
    "Robert Axelrod (Complejidad de la Cooperación)": {
        "desc": "Teoría de Juegos, Evolución de la Cooperación y Normas.",
        "preguntas": [
            "El Dilema del Prisionero: ¿Existen incentivos estructurales que hacen racional la traición individual?",
            "Estrategia Tit-for-Tat: ¿Está el actor respondiendo con reciprocidad estricta? ¿Está respondiendo proporcionalmente o escalando?",
            "La Sombra del Futuro: ¿Es la interacción lo suficientemente duradera para fomentar la cooperación? ¿Tienen expectativas de interactuar nuevamente?",
            "Meta-Normas: ¿Existe presión social o sanciones de terceros para castigar a los desertores?",
            "Detección de Trampas: ¿Qué mecanismos de verificación existen para asegurar el cumplimiento?",
            "Estructura de Pagos: ¿Cómo alterar los incentivos para que cooperar sea más rentable que traicionar?"
        ]
    },
    "Teoría de Juegos (John Nash)": {
        "desc": "Equilibrios matemáticos en la toma de decisiones.",
        "preguntas": [
            "Suma Cero vs Suma Variable: ¿Para que uno gane, el otro debe perderlo todo?",
            "Equilibrio de Nash: ¿Cuál es la situación donde nadie tiene incentivos para cambiar su estrategia?",
            "La Gallina (Chicken Game): ¿Quién cederá primero ante la inminencia del choque?"
        ]
    },

    # -------------------------------------------------------------------------
    # 🧠 BLOQUE 6: TOMA DE DECISIONES Y ANÁLISIS ESTRATÉGICO
    # -------------------------------------------------------------------------
    "--- TOMA DE DECISIONES Y SEGURIDAD ---": { "desc": "", "preguntas": [] },

    "Graham Allison (Los 3 Modelos de Decisión)": {
        "desc": "Análisis de la crisis desde múltiples lentes (La Esencia de la Decisión).",
        "preguntas": [
            "Modelo I (Actor Racional): ¿Cuál es la opción lógica que maximiza beneficios y minimiza costos estratégicos?",
            "Modelo II (Proceso Organizacional): ¿Qué procedimientos estándar (SOPs) y rutinas limitan la flexibilidad del gobierno?",
            "Modelo III (Política Burocrática): ¿Qué agencias o individuos internos están luchando por el poder y cómo afecta esto la decisión final?"
        ]
    },
    "Barry Buzan (Seguridad Integral y Securitización)": {
        "desc": "Los 5 Sectores de Seguridad y la Teoría de la Securitización.",
        "preguntas": [
            "Análisis Multisectorial: Evalúa amenazas en los 5 sectores: Militar, Político, Económico, Societal y Ambiental.",
            "Nivel Sistémico: ¿Cómo influye la anarquía internacional o la polaridad en el conflicto?",
            "Nivel Estatal: ¿Qué presiones burocráticas o nacionales limitan al Estado?",
            "Nivel Individual: ¿El perfil psicológico de los líderes altera la toma de decisiones?",
            "Seguridad Societal: ¿Está amenazada la identidad colectiva (religión, etnia, cultura)?",
            "Actor Securitizador: ¿Quién está declarando el asunto como una 'amenaza existencial'?",
            "Objeto Referente: ¿Qué es exactamente lo que se intenta proteger (El Estado, la Nación, la Economía)?",
            "Medidas Extraordinarias: ¿Se está usando la retórica de seguridad para justificar acciones fuera de la política normal?"
        ]
    },
    "John Boyd (Ciclo OODA)": {
        "desc": "Velocidad de decisión en conflicto (Observar, Orientar, Decidir, Actuar).",
        "preguntas": [
            "Velocidad del Ciclo: ¿Quién está completando su ciclo OODA más rápido?",
            "Fase de Orientación: ¿Cómo los sesgos culturales y la herencia genética moldean la percepción del adversario?",
            "Colapso del Adversario: ¿Cómo podemos generar ambigüedad para aislar al enemigo de su entorno?"
        ]
    },
    "Sherman Kent (Doctrina de Inteligencia Estratégica)": {
        "desc": "Los fundamentos clásicos: Inteligencia como Conocimiento, Organización y Actividad.",
        "preguntas": [
            "La Pirámide de Kent: Clasifica la información analizada. ¿Es Nivel 1 (Hechos/Descriptivo), Nivel 2 (Tendencias/Explicativo) o Nivel 3 (Estimativo/Predictivo)?",
            "Lenguaje Probabilístico: ¿Se utilizan términos de probabilidad estimativa precisos (ej: 'Muy Probable', 'Posibilidad Remota') o se usa lenguaje ambiguo para evitar responsabilidad (ej: 'podría', 'quizás')?",
            "Ecuación de Riesgo: ¿Se están evaluando las 'Capacidades' (lo que el adversario PUEDE hacer) separadas de las 'Intenciones' (lo que QUIERE hacer)?",
            "Relación Productor-Consumidor: ¿El análisis mantiene la objetividad necesaria o parece haber sido redactado para complacer una decisión política ya tomada (Politización)?"
        ]
    },
    "Marcelo de los Reyes (Inteligencia y RRII)": {
        "desc": "La inteligencia como insumo crítico para la política exterior.",
        "preguntas": [
            "Politización de la Inteligencia: ¿Se está produciendo inteligencia para complacer al decisor político (Inteligencia a la carta)?",
            "Diplomacia Paralela: ¿Están los servicios de inteligencia actuando como canales diplomáticos secretos?",
            "Sorpresa Estratégica: ¿Falló la inteligencia por falta de datos o por falta de imaginación para interpretar las señales?"
        ]
    },
    
    # =========================================================================
    # 🌐 BLOQUE 7: CIBERINTELIGENCIA Y REDES
    # =========================================================================
    "--- CIBERESPACIO E INFO ---": { "desc": "", "preguntas": [] },

    "Cyber Kill Chain (Lockheed Martin)": {
        "desc": "Fases de una intrusión cibernética.",
        "preguntas": [
            "Reconocimiento: ¿Qué datos se están recolectando antes del ataque?",
            "Armamentización: ¿Cómo se creó el malware o el exploit?",
            "Entrega y Explotación: ¿Fue phishing, USB, vulnerabilidad web?",
            "Acciones sobre Objetivos: ¿Se busca robar datos, destruir sistemas o secuestrar (Ransomware)?"
        ]
    },
    "Teoría del Actor-Red (Latour)": {
        "desc": "Humanos y objetos (algoritmos) tienen agencia.",
        "preguntas": [
            "Agencia Tecnológica: ¿Cómo un algoritmo o plataforma está moldeando el conflicto por sí solo?",
            "Cajas Negras: ¿Qué procesos técnicos se están aceptando sin cuestionar su funcionamiento?",
            "Traducción: ¿Cómo se están redefiniendo los intereses a través de la red?"
        ]
    },
    "Modelo Diamante de Intrusión": {
        "desc": "Relación entre Adversario, Infraestructura, Capacidad y Víctima.",
        "preguntas": [
            "Eje Adversario-Víctima: ¿Cuál es la intención sociopolítica detrás del ataque técnico?",
            "Eje Infraestructura-Capacidad: ¿Qué servidores o IPs (Infraestructura) soportan el malware (Capacidad)?",
            "Pivoteo: ¿Podemos usar la infraestructura detectada para encontrar otras víctimas desconocidas?"
        ]
    },

    # =========================================================================
    # 🧠 BLOQUE 8: PSICOLOGÍA Y MENTE DEL ADVERSARIO
    # =========================================================================
    "--- PSICOLOGÍA OPERATIVA ---": { "desc": "", "preguntas": [] },

    "Robert M. Ryder (Conciencia de Dominio / Domain Awareness)": {
        "desc": "Comprensión holística y cognitiva del entorno operativo total.",
        "preguntas": [
            "Ceguera de Dominio: ¿Qué esfera del entorno (marítima, ciber, espacial, humana) estamos ignorando por falta de sensores?",
            "Fusión de Datos: ¿Se están conectando puntos aislados para formar una imagen operativa común (COP)?",
            "Anticipación Cognitiva: ¿Estamos reaccionando a eventos o previendo flujos en el entorno?",
            "Conciencia Cultural: ¿Entendemos el 'terreno humano' tan bien como el terreno físico?"
        ]
    },
    "Perfilado Dark Triad (Tríada Oscura)": {
        "desc": "Psicopatía, Narcisismo y Maquiavelismo en el liderazgo.",
        "preguntas": [
            "Narcisismo: ¿El líder necesita admiración constante y reacciona con ira a la crítica?",
            "Maquiavelismo: ¿Manipula a aliados y enemigos sin remordimiento?",
            "Psicopatía: ¿Muestra falta total de empatía y toma riesgos impulsivos?",
            "Vulnerabilidad del Ego: ¿Cómo se puede explotar su necesidad de validación?"
        ]
    },
    "Código MICE (Motivaciones de Traición)": {
        "desc": "Money, Ideology, Coercion, Ego.",
        "preguntas": [
            "Dinero (Money): ¿Existen crisis financieras personales?",
            "Ideología (Ideology): ¿Cree el sujeto en una causa superior opuesta?",
            "Coerción (Coercion): ¿Existe material de chantaje (Kompromat)?",
            "Ego: ¿Se siente infravalorado o busca venganza?"
        ]
    },
    "Allan Pease (Lenguaje No Verbal y Poder)": {
        "desc": "Lectura de gestos, posturas y congruencia.",
        "preguntas": [
            "Congruencia: ¿Lo que dice el líder verbalmente coincide con sus gestos? (Si no, el gesto dice la verdad).",
            "Gestos de Poder y Dominio: ¿Usa la 'Cúpula de Poder' (manos), palmas ocultas o toma de espacio territorial?",
            "Microexpresiones: ¿Hay fugas faciales de miedo, desprecio o ira en momentos clave del discurso?"
        ]
    },
    "Gustave Le Bon (Psicología de Masas)": {
        "desc": "Comportamiento irracional y contagio emocional.",
        "preguntas": [
            "Contagio Mental: ¿Cómo se propaga la emoción irracional?",
            "Líder de Masas: ¿Quién canaliza el odio o la esperanza de la multitud?",
            "Imágenes Simplistas: ¿Qué eslóganes reemplazan el pensamiento lógico?"
        ]
    },
    "David Alandete (Fake News: Arma de Destrucción Masiva)": {
        "desc": "Desinformación, polarización y algoritmos.",
        "preguntas": [
            "Objetivo de la Desinformación: ¿Buscan convencer de una mentira o simplemente sembrar duda y caos para paralizar la sociedad?",
            "Cámaras de Eco: ¿Se está usando el algoritmo para radicalizar grupos específicos?",
            "Actores Proxy: ¿Quién está amplificando el mensaje? (Bots, tontos útiles, medios estatales disfrazados)."
        ]
    },
    "Edward Bernays (Propaganda y Relaciones Públicas)": {
        "desc": "La ingeniería del consentimiento y la manipulación de la 'mente de grupo'.",
        "preguntas": [
            "Autoridad de Terceros: ¿El mensaje utiliza a 'expertos independientes' (médicos, científicos) para vender una idea y burlar el escepticismo?",
            "Creación de Eventos: ¿La noticia es un hecho espontáneo o un 'pseudo-evento' fabricado para generar cobertura mediática?",
            "Apelación al Instinto Gregario: ¿Se presenta la idea como algo que 'todos los líderes inteligentes' ya aceptan, aislando al disidente?"
        ]
    },
    "Cass R. Sunstein (Rumorología y Cascadas de Información)": {
        "desc": "Cómo se propagan las creencias falsas y la polarización de grupos.",
        "preguntas": [
            "Cascadas de Disponibilidad: ¿El rumor se ha vuelto creíble solo porque 'se repite en todas partes' (reputación social) y no por evidencia real?",
            "Polarización de Grupo: ¿El debate interno está volviendo al grupo más extremo de lo que era al principio?",
            "Asimilación Sesgada: ¿Los actores están aceptando ciegamente la información que confirma sus sesgos y descartando agresivamente la evidencia contraria?"
        ]
    },
    "Wilson Bryan Key (Seducción Subliminal)": {
        "desc": "Estímulos ocultos que apelan al subconsciente (Eros y Thanatos).",
        "preguntas": [
            "Implantes Emocionales: ¿Existen elementos visuales o auditivos ocultos diseñados para provocar ansiedad o deseo sexual sin pasar por el filtro racional?",
            "Sobrecarga Sensorial: ¿Se está bombardeando al consciente con datos irrelevantes para que el subconsciente acepte una orden oculta?",
            "Apelación a la Muerte/Deseo: ¿El discurso o imagen juega con los miedos primarios a la extinción o con los impulsos reproductivos?"
        ]
    },
    "Harold Lasswell (Teoría de la Propaganda)": {
        "desc": "Gestión de actitudes colectivas y símbolos.",
        "preguntas": [
            "Fórmula de Lasswell: ¿QUIÉN (control) dice QUÉ (contenido) a QUIÉN (audiencia) en qué CANAL (medio) con qué EFECTO?",
            "Gestión de Símbolos: ¿Qué mitos, himnos o banderas se están manipulando para evocar emociones irracionales?",
            "Movilización del Odio: ¿Se está dirigiendo la agresividad colectiva hacia un enemigo común fabricado?"
        ]
    },
    "Joseph Goebbels (Matriz de Propaganda Completa)": {
        "desc": "Auditoría forense de los 11 principios de manipulación de masas.",
        "preguntas": [
            "1. Principio de Simplificación y Enemigo Único: ¿Se reduce toda la complejidad del problema a un solo símbolo, eslogan o enemigo a batir?",
            "2. Principio del Método de Contagio: ¿Se reúnen diversos adversarios (que no tienen nada que ver entre sí) bajo una sola categoría o etiqueta negativa para demonizarlos en bloque?",
            "3. Principio de la Transposición: ¿Carga el emisor sus propios errores o defectos sobre el adversario? ('Acusa al otro de lo que tú haces').",
            "4. Principio de la Exageración y Desfiguración: ¿Se convierte un hecho anecdótico, pequeño o aislado en una amenaza grave o crisis sistémica?",
            "5. Principio de la Vulgarización: ¿El mensaje está adaptado deliberadamente al nivel menos inteligente de la audiencia, evitando argumentos racionales complejos?",
            "6. Principio de Orquestación: ¿Se limitan a pocas ideas básicas pero las repiten incansablemente desde diferentes ángulos o perspectivas?",
            "7. Principio de Renovación: ¿Se emiten nuevas acusaciones o informaciones tan rápido que, cuando el adversario responde, el público ya está interesado en otra cosa?",
            "8. Principio de la Verosimilitud: ¿Se construyen argumentos basándose en fuentes parciales, fragmentos de verdad o 'globos sonda' para dar credibilidad a una mentira mayor?",
            "9. Principio de la Silenciación: ¿Se omiten o acallan sistemáticamente las noticias que favorecen al rival o las que contradicen la narrativa oficial?",
            "10. Principio de la Transfusión: ¿Se opera sobre un sustrato de odio, prejuicio, mitología o tradición preexistente en la cultura para potenciar el mensaje?",
            "11. Principio de la Unanimidad: ¿Se crea la falsa impresión de que 'todo el mundo' piensa así y que el disidente está socialmente aislado?"
        ]
    },
        
    # =========================================================================
    # 🔮 BLOQUE 9: PROSPECTIVA Y COMPLEJIDAD
    # =========================================================================
    "--- FUTUROS Y SISTEMAS ---": { "desc": "", "preguntas": [] },

    "Análisis Causal por Capas (CLA - Inayatullah)": {
        "desc": "Deconstrucción profunda de la realidad.",
        "preguntas": [
            "La Letanía: ¿Qué dicen los titulares oficiales?",
            "Causas Sistémicas: ¿Qué estructuras generan el problema?",
            "Visión del Mundo: ¿Qué ideologías sostienen el sistema?",
            "Mito y Metáfora: ¿Cuál es la historia inconsciente detrás de todo?"
        ]
    },
    "Michel Godet (Prospectiva Estratégica)": {
        "desc": "De la anticipación a la acción (El Triángulo Griego: Anticipación, Apropiación, Acción).",
        "preguntas": [
            "Preactividad vs Proactividad: ¿Nos estamos preparando para cambios esperados (seguro) o estamos provocando los cambios deseados (apuesta)?",
            "Actitud Estratégica: ¿El actor es reactivo (apaga fuegos), preactivo (se prepara para lo inevitable) o proactivo (provoca el cambio deseado)?",
            "Variables Clave (MICMAC): Identifica las variables 'motrices' ocultas que controlan el sistema (causas raíz) vs las variables 'dependientes' (síntomas).",
            "Juego de Actores (MACTOR): ¿Qué convergencias (aliados) y divergencias (conflictos) de objetivos existen y quién tiene la fuerza para imponer su voluntad?",
            "El Camino Estratégico: Contrasta el 'Escenario Probable' (tendencial) con el 'Escenario Deseable'. ¿Qué acciones concretas deben tomarse para cerrar esa brecha?"
        ]
    },
    "Nassim Taleb (Cisne Negro & Antifragilidad)": {
        "desc": "Gestión de lo improbable y el caos.",
        "preguntas": [
            "Cisne Negro: Evento de probabilidad baja e impacto infinito.",
            "Rinoceronte Gris: Amenaza obvia ignorada voluntariamente.",
            "Antifragilidad: ¿Qué actor se beneficia del desorden?"
        ]
    },
    "Análisis de Señales Débiles (Weak Signals)": {
        "desc": "Detección temprana de anomalías.",
        "preguntas": [
            "Ruido Marginal: ¿Qué dato 'irrelevante' se repite?",
            "Ceguera Experta: Identifica qué escenarios están siendo descartados por los expertos oficiales por considerarlos 'imposibles' o 'ridículos'.",
            "Patrones de Rareza: ¿Qué evento rompe la continuidad histórica?"
        ]
    },

    # =========================================================================
    # 🕵️ BLOQUE 10: VERIFICACIÓN
    # =========================================================================
    "--- ANÁLISIS COMPARATIVO ---": { "desc": "", "preguntas": [] },

    "Triangulación de Fuentes (Cross-Check)": {
        "desc": "Técnica forense para detectar contradicciones, mentiras y vacíos entre múltiples documentos.",
        "preguntas": [
            "Matriz de Contradicciones: Genera una tabla comparativa identificando EXCLUSIVAMENTE los puntos donde el 'Documento A' dice algo diferente al 'Documento B' (Fechas, cifras, nombres).",
            "Análisis de Silencios: ¿Qué información crucial menciona una fuente pero es omitida sospechosamente por las otras?",
            "Divergencia Narrativa: ¿Cómo cambia el tono o la intención política entre una versión y otra?",
            "Veredicto de Credibilidad: Basado en la consistencia interna y externa, ¿qué fuente parece tener mayor acceso a la verdad y cuál parece intoxicada?"
        ]
    },
    "Análisis de Decepción y Engaño": {
        "desc": "Detección de manipulación informativa.",
        "preguntas": [
            "Señuelos: ¿Existe información demasiado perfecta o detallada diseñada para atraer nuestra atención lejos de lo importante?",
            "Canales de Retroalimentación: ¿El adversario nos está diciendo lo que queremos escuchar (Sesgo de confirmación)?",
            "Inconsistencias Temporales: ¿Hay eventos reportados en una secuencia cronológica imposible?"
        ]
    },
    
    # =========================================================================
    # 🛠️ BLOQUE 11: HERRAMIENTAS TÁCTICAS (SATs)
    # =========================================================================
    "--- HERRAMIENTAS ESTRUCTURADAS ---": { "desc": "", "preguntas": [] },

    "Análisis de Hipótesis en Competencia (ACH)": { "desc": "Matriz científica para evitar sesgos.", "preguntas": ["Generación de Hipótesis.", "Matriz de Evidencia.", "Diagnóstico de Consistencia.", "Refutación."] },
    "Análisis de Actores (Stakeholder Mapping)": { "desc": "Mapa de poder e intereses.", "preguntas": ["Matriz Poder/Interés.", "Vetadores.", "Spoilers (Saboteadores)."] },
    "Matriz CARVER (Selección de Objetivos)": { "desc": "Evaluación de blancos.", "preguntas": ["Criticidad.", "Accesibilidad.", "Recuperabilidad.", "Vulnerabilidad.", "Efecto.", "Reconocibilidad."] },
    "Análisis PMESII-PT (Entorno Operativo)": { "desc": "Análisis holístico.", "preguntas": ["Político/Militar.", "Económico/Social.", "Información/Infraestructura.", "Físico/Tiempo."] },
    "Análisis FODA (SWOT) de Inteligencia": { "desc": "Ofensivo/Defensivo.", "preguntas": ["Amenazas Inminentes.", "Oportunidades.", "Vulnerabilidades Internas.", "Fortalezas."] },
    "Técnica de los 5 Porqués": { "desc": "Búsqueda de Causa Raíz.", "preguntas": ["Síntoma.", "¿Por qué? (x5).", "Falla Sistémica."] },
    "Abogado del Diablo": { "desc": "Desafío de asunciones.", "preguntas": ["Desafío Frontal a la tesis principal.", "Defensa de la postura irracional del adversario."] },
    "Richards J. Heuer (Psicología del Análisis de Inteligencia)": { "desc": "Chequeo de sesgos cognitivos del propio analista.", "preguntas": ["Sesgo de Confirmación: ¿Estamos buscando solo información que confirma nuestra hipótesis y descartando la que la contradice?", "Imagen en Espejo: ¿Estamos asumiendo que el adversario piensa y actúa racionalmente como nosotros?", "Anclaje: ¿Estamos demasiado atados a la primera estimación o dato que recibimos al inicio de la crisis?"
         ]
    }
}

# ==========================================
# 📘 TEXTO DEL MANUAL (CONTENIDO ESTÁTICO)
# ==========================================
MANUAL_USUARIO = """
# 📘 MANUAL INTEGRAL | SISTEMA STRATINTEL SOLUTIONS

## PARTE 1: OPERACIONES TÉCNICAS
**1. INTRODUCCIÓN**
StratIntel es un Sistema de Soporte a la Decisión (DSS) que utiliza IA para aplicar marcos teóricos de inteligencia y relaciones internacionales a documentos no estructurados.

**2. FLUJO DE TRABAJO BÁSICO**
1.  **Ingesta:** Suba sus PDFs, DOCXs o pegue texto en la pestaña correspondiente.
2.  **Configuración:** Ingrese su API Key en el menú lateral.
3.  **Selección:** Elija el Marco Teórico adecuado para su misión (Ver Parte 2).
4.  **Profundidad:**
    * *Estratégico:* Resumen ejecutivo (BLUF).
    * *Táctico:* Responde todas las preguntas teóricas.
    * *Operacional:* Selección manual de preguntas.

**3. HERRAMIENTAS ESPECIALES**
* **🎨 Visualización:** (Si está activo) Genera esquemas de actores al final del reporte.
* **🕵️ Contrainteligencia:** Cargue 2+ documentos y use la técnica "Triangulación" para hallar contradicciones.
* **💾 Exportación:** Use los botones al final para descargar en Word o PDF.

---

## PARTE 2: DICCIONARIO DE DOCTRINA Y ANÁLISIS
*Guía de referencia para interpretar los resultados de cada técnica.*

### 🌍 BLOQUE 1: ESTRUCTURA Y PODER
* **Fuerzas Profundas (Renouvin & Duroselle):** Analiza si el conflicto es causado por la voluntad de un líder o por corrientes históricas inevitables (geografía, demografía).
* **Ciclos Imperiales (Kennedy & Duroselle):** Busca síntomas de "Sobrestiramiento Imperial" (gasto militar excesivo) o decadencia interna en grandes potencias.
* **Venganza de la Geografía (Kaplan):** Identifica cómo el mapa físico (montañas, ríos) dicta el destino político y las zonas de conflicto (Shatterbelts).
* **Realismo Clásico (Morgenthau, Carr, Cline):**
    * *Morgenthau:* Evalúa el interés nacional en términos de poder y supervivencia, ignorando la moral.
    * *Carr:* Detecta la hipocresía de potencias que disfrazan intereses egoístas como "valores universales".
    * *Cline:* Calcula el potencial de poder: $Pp = (C + E + M) * (S + W)$.
* **Geopolítica Dura (Mackinder & Spykman):** Analiza el control territorial global. ¿Quién domina el "Corazón Continental" (Heartland) y quién los bordes costeros (Rimland)?
* **Neorrealismo (Waltz & Mearsheimer):**
    * *Defensivo (Waltz):* El estado solo busca seguridad.
    * *Ofensivo (Mearsheimer):* El estado busca hegemonía total y aprovechará cualquier oportunidad para debilitar rivales.

### 🤝 BLOQUE 2: INSTITUCIONES E IDENTIDAD
* **Poder Multidimensional (Nye):** Evalúa el uso de *Soft Power* (atracción cultural) y *Smart Power* (combinación de fuerza y diplomacia).
* **Paz Liberal (Kant & Keohane):** Analiza si el comercio y las instituciones internacionales hacen que la guerra sea "demasiado costosa" (Interdependencia).
* **Constructivismo (Wendt & Huntington):**
    * *Wendt:* ¿El enemigo es una construcción social?
    * *Huntington:* ¿Es un choque de civilizaciones (identidad cultural/religiosa) y no ideológico?

### ⚔️ BLOQUE 3: ESTRATEGIA MILITAR
* **Aproximación Indirecta (Liddell Hart):** Verifica si el actor ataca la debilidad del enemigo evitando el choque frontal (físico o psicológico).
* **Nuevas Guerras (Kaldor & Creveld):** Para conflictos modernos donde se mezcla crimen, limpieza étnica y actores no estatales (cárteles, terrorismo).
* **Guerra Irrestricta (Qiao Liang):** Identifica el uso de "todo" como arma: leyes (Lawfare), economía, drogas y medios de comunicación.

### 💰 BLOQUE 4: GEOECONOMÍA Y CAOS
* **Economía Ilícita (Naím):** Analiza el poder de redes criminales (narco, tráfico) que erosionan al Estado.
* **Geoeconomía (Luttwak):** El uso de la gramática del comercio (sanciones, deuda, inversiones predatorias) para fines de guerra.

### 🤝 BLOQUE 5: NEGOCIACIÓN Y JUEGOS
* **Estrategia del Conflicto (Schelling):** Analiza el uso racional de la irracionalidad, la disuasión y los puntos focales en una crisis.
* **Teoría de Juegos (Axelrod & Nash):** Modela matemáticamente la cooperación. ¿Estamos en un dilema del prisionero (traición incentiva) o juego de la gallina (choque inminente)?

### 🧠 BLOQUE 6: TOMA DE DECISIONES
* **Modelos de Decisión (Allison):** Determina si una acción fue racional (Modelo I), una rutina burocrática (Modelo II) o resultado de peleas políticas internas (Modelo III).
* **Ciclo OODA (Boyd):** Evalúa la velocidad de reacción: Observar, Orientar, Decidir, Actuar. Quien completa el ciclo más rápido, gana.
* **Sherman Kent (Calidad de Inteligencia):** **AUDITORÍA.** Detecta lenguaje ambiguo, confusión entre Capacidad e Intención, y politización del informe.

### 🌐 BLOQUE 7: CIBERINTELIGENCIA
* **Cyber Kill Chain:** Desglosa ataques digitales en fases (Reconocimiento -> Entrega -> Explotación).
* **Modelo Diamante:** Relaciona al Adversario con su Infraestructura, Capacidades y Víctimas.

### 🧠 BLOQUE 8: PSICOLOGÍA Y MENTE
* **Perfilado Oscuro (Dark Triad & MICE):** Evalúa líderes por Narcisismo/Psicopatía y busca motivos de traición (Dinero, Ideología, Coerción, Ego).
* **Propaganda (Goebbels, Bernays, Sunstein):**
    * *Goebbels:* Principios de simplificación y repetición.
    * *Bernays:* Ingeniería del consentimiento mediante "autoridad de terceros".
    * *Sunstein:* Viralidad de rumores y cámaras de eco.
    * *Key:* Seducción subliminal y ataques al subconsciente.

### 🔮 BLOQUE 9: PROSPECTIVA (FUTUROS)
* **Michel Godet:** Construcción de escenarios. Distingue variables motrices (causas) de dependientes (síntomas) y define el juego de actores.
* **Cisnes Negros (Taleb):** Identifica eventos improbables de alto impacto o amenazas obvias ignoradas (Rinocerontes Grises).

### 🕵️ BLOQUE 10: VERIFICACIÓN
* **Triangulación (Cross-Check):** Técnica forense. Compara documentos para hallar contradicciones, silencios y cambios de narrativa.

---
*Documentación Oficial del Sistema StratIntel Solutions | Uso Reservado*
"""

# --- GESTIÓN DE ESTADO ---
if 'api_key' not in st.session_state: st.session_state['api_key'] = ""
if 'texto_analisis' not in st.session_state: st.session_state['texto_analisis'] = ""
if 'origen_dato' not in st.session_state: st.session_state['origen_dato'] = "Ninguno"

# --- FUNCIONES DE PROCESAMIENTO ---
def buscar_en_web(query):
    try:
        search = DuckDuckGoSearchRun()
        return search.run(query)
    except Exception as e: return f"Error web: {e}"

def procesar_archivos_pdf(archivos):
    texto_total = ""
    nombres = []
    for archivo in archivos:
        reader = pypdf.PdfReader(archivo)
        texto_pdf = "".join([p.extract_text() for p in reader.pages])
        texto_total += f"\n--- ARCHIVO: {archivo.name} ---\n{texto_pdf}\n"
        nombres.append(archivo.name)
    return texto_total, str(nombres)

def procesar_archivos_docx(archivos):
    texto_total = ""
    nombres = []
    for archivo in archivos:
        doc = Document(archivo)
        texto_doc = "\n".join([para.text for para in doc.paragraphs])
        texto_total += f"\n--- ARCHIVO: {archivo.name} ---\n{texto_doc}\n"
        nombres.append(archivo.name)
    return texto_total, str(nombres)

def obtener_texto_web(url):
    try:
        h = {'User-Agent': 'Mozilla/5.0'}
        r = requests.get(url, headers=h, timeout=15)
        s = BeautifulSoup(r.content, 'html.parser')
        for script in s(["script", "style"]): script.extract()
        return s.get_text(separator='\n')
    except Exception as e: return f"Error: {e}"
   
def generar_esquema_graphviz(texto_analisis, api_key):
    """Genera código DOT con código de colores semántico y TÍTULO dinámico."""
    try:
        if not api_key: return None, "Falta API Key de Google para visualización."
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")
        
        prompt = f"""
        ACTÚA COMO: Experto en Visualización de Inteligencia (Link Analysis).
        OBJETIVO: Generar un GRAFO (DOT Graphviz) que resuma las relaciones clave y tenga un TÍTULO DESCRIPTIVO.
        
        REGLAS DE COLOR OBLIGATORIAS:
        1. 🟧 ACTORES (Países, Líderes, Organizaciones): fillcolor="#ffcc99" (Naranja)
        2. 🟥 AMENAZAS (Conflictos, Riesgos, Crisis, Ataques): fillcolor="#ffcccc" (Rojo Claro)
        3. 🟦 CONCEPTOS (Teorías, Doctrinas, Economía, Recursos, Intereses): fillcolor="#ccddff" (Azul Claro)
        
        INSTRUCCIONES TÉCNICAS:
        1. Analiza el texto e identifica las entidades y relaciones más críticas.
        2. Genera un TÍTULO CORTO y conciso (máximo 8 palabras) que resuma el tema principal del análisis.
        3. Inserta el título al inicio del grafo usando el atributo 'label'.
        4. Genera SOLO el código DOT válido. Sin markdown.
        
        EJEMPLO DE ESTRUCTURA ESPERADA:
        digraph G {{
            # --- CONFIGURACIÓN DEL TÍTULO ---
            graph [label="TÍTULO GENERADO POR LA IA AQUÍ", labelloc=t, fontsize=16, fontname="Arial Bold", fontcolor="#333333"];
            rankdir=LR;
            node [style=filled, fontname="Arial", shape=box];
            edge [fontname="Arial", fontsize=10];
            
            # --- NODOS Y RELACIONES ---
            "EEUU" [fillcolor="#ffcc99", label="Actor: EEUU"];
            "Guerra Híbrida" [fillcolor="#ffcccc", label="Amenaza: Guerra Híbrida"];
            "Realismo Ofensivo" [fillcolor="#ccddff", label="Concepto: Realismo Ofensivo"];
            "EEUU" -> "Guerra Híbrida" [label="enfrenta riesgo de"];
            "Guerra Híbrida" -> "Realismo Ofensivo" [label="analizada vía"];
        }}
        
        TEXTO A ANALIZAR:
        {texto_analisis[:15000]}
        """
        
       res = model.generate_content(prompt)
        codigo_dot = res.text.replace("```dot", "").replace("```", "").replace("DOT", "").strip()
        grafico = graphviz.Source(codigo_dot)
        return grafico, None
    except Exception as e:
        return None, f"Error visual: {e}"

# --- FUNCIONES DE REPORTE ---
def limpiar_texto(t):
    if not t: return ""
    reps = {"✨": "", "🚀": "", "⚠️": "[!]", "✅": "[OK]", "🛡️": "", "🔒": "", "🎖️": "", "♟️": "", "⚖️": ""}
    for k,v in reps.items(): t = t.replace(k,v)
    return t.encode('latin-1', 'replace').decode('latin-1')

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'StratIntel Report', 0, 1, 'C')
        self.ln(5)
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 7)
        self.cell(0, 10, 'Generado por IA. Uso Reservado.', 0, 0, 'C')

def crear_pdf(texto, tecnicas, fuente):
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_font("Arial", "B", 12)
    pdf.multi_cell(0, 7, limpiar_texto(f"Fuente: {fuente}\nTécnicas: {tecnicas}"))
    pdf.ln(5)
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 7, limpiar_texto(texto))
    return pdf.output(dest='S').encode('latin-1', 'replace')

def crear_word(texto, tecnicas, fuente):
    doc = Document()
    doc.add_heading('StratIntel Intelligence Report', 0)
    doc.add_paragraph(f"Fuente: {fuente}").bold = True
    doc.add_paragraph(f"Técnicas: {tecnicas}").bold = True
    for l in texto.split('\n'):
        if l.startswith('#'): doc.add_heading(l.replace('#','').strip(), level=2)
        else: doc.add_paragraph(l)
    aviso = doc.add_paragraph()
    aviso.add_run("\n\n------------------\nAVISO: Generado por IA.").font.size = 8
    b = BytesIO(); doc.save(b); b.seek(0)
    return b

# ==========================================
# 🖥️ INTERFAZ PRINCIPAL
# ==========================================
st.sidebar.title("♟️ StratIntel OS")
st.sidebar.caption("v2.0 | Hybrid Engine")
st.sidebar.markdown("---")

# CARGA SEGURA DE CLAVES (Sin crashear si faltan)
google_key_auto = st.secrets.get("GOOGLE_API_KEY", "")
router_key_auto = st.secrets.get("OPENROUTER_API_KEY", "")

# SELECTOR DE MISION
st.sidebar.subheader("🎯 Misión")
tecnicas_seleccionadas = st.sidebar.multiselect(
    "Técnicas (Máx 3):",
    options=list(DB_CONOCIMIENTO.keys()),
    max_selections=3
)
temp = st.sidebar.slider("Creatividad", 0.0, 1.0, 0.4)
if st.sidebar.button("🔒 Salir"): del st.session_state["password_correct"]; st.rerun()

st.title("♟️ StratIntel Solutions | División de Análisis")

# PESTAÑAS DE CARGA
t1, t2, t3, t4 = st.tabs(["📂 PDFs", "📝 DOCXs", "🌐 Web", "✍️ Manual"])
with t1:
    f = st.file_uploader("PDFs", type="pdf", accept_multiple_files=True)
    if f and st.button("Procesar PDF"):
        t, n = procesar_archivos_pdf(f); st.session_state['texto_analisis']=t; st.session_state['origen_dato']=f"PDFs: {n}"; st.success(f"✅ {len(f)}")
with t2:
    f = st.file_uploader("DOCXs", type="docx", accept_multiple_files=True)
    if f and st.button("Procesar DOCX"):
        t, n = procesar_archivos_docx(f); st.session_state['texto_analisis']=t; st.session_state['origen_dato']=f"DOCXs: {n}"; st.success(f"✅ {len(f)}")
with t3:
    u = st.text_input("URL")
    if st.button("Web"): st.session_state['texto_analisis']=obtener_texto_web(u); st.session_state['origen_dato']=f"Web: {u}"; st.success("OK")
with t4:
    m = st.text_area("Texto Manual")
    if st.button("Fijar Texto"): st.session_state['texto_analisis']=m; st.session_state['origen_dato']="Manual"; st.success("OK")

st.markdown("---")
if st.session_state['texto_analisis']:
    with st.expander(f"Fuente Activa: {st.session_state['origen_dato']}"): st.write(st.session_state['texto_analisis'][:1000])

# ==========================================
# 🚀 EJECUCIÓN HÍBRIDA (EL CEREBRO)
# ==========================================
st.header("Generación de Informe")

if not st.session_state['texto_analisis']:
    st.warning("⚠️ Carga datos para comenzar.")
else:
    c1, c2 = st.columns([1, 2])
    with c1:
        if not tecnicas_seleccionadas: st.info("👈 Selecciona técnicas.")
        
        profundidad = st.radio(
            "Profundidad:", 
            ["🔍 Estratégico", "🎯 Táctico", "⚙️ Operacional"],
            help="Estratégico: Resumen. Táctico: Completo. Operacional: Manual."
        )
        
        # Selección Manual
        preguntas_manuales = {}
        if "Operacional" in profundidad and tecnicas_seleccionadas:
            for tec in tecnicas_seleccionadas:
                qs = DB_CONOCIMIENTO.get(tec, {}).get("preguntas", [])
                if qs:
                    sel = st.multiselect(f"Q: {tec}:", qs)
                    preguntas_manuales[tec] = sel
        
        usar_internet = st.checkbox("🌐 Búsqueda Web")
        pir = st.text_area("PIR (Opcional):", height=80)

    with c2:
        # --- SELECTOR DE MOTOR ---
        st.markdown("### 🧠 Motor de IA")
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            PROVEEDOR = st.radio("Proveedor:", ["Google Gemini", "DeepSeek (OpenRouter)"], label_visibility="collapsed")
        with col_m2:
            # Lógica de Claves Inteligente
            api_key_final = ""
            if "Google" in PROVEEDOR:
                if google_key_auto:
                    st.success("🔑 Google Key Detectada")
                    api_key_final = google_key_auto
                else:
                    api_key_final = st.text_input("🔑 Pega tu Google Key:", type="password")
            else:
                if router_key_auto:
                    st.success("🔑 OpenRouter Key Detectada")
                    api_key_final = router_key_auto
                else:
                    api_key_final = st.text_input("🔑 Pega OpenRouter Key:", type="password")

        # BOTÓN EJECUTAR
        if st.button("🚀 EJECUTAR MISIÓN", type="primary", use_container_width=True, disabled=len(tecnicas_seleccionadas)==0):
            if not api_key_final:
                st.error("❌ Falta la API Key para el motor seleccionado.")
            else:
                try:
                    # Limpieza
                    if 'codigo_dot_mapa' in st.session_state: del st.session_state['codigo_dot_mapa']
                    if 'res' in st.session_state: del st.session_state['res']
                    
                    # Guardamos la key de Google en sesión SIEMPRE, porque se usa para el mapa visual
                    if "Google" in PROVEEDOR: st.session_state['api_key'] = api_key_final
                    elif google_key_auto: st.session_state['api_key'] = google_key_auto

                    # Configura Gemini por si acaso (para mapas)
                    if st.session_state.get('api_key'):
                        genai.configure(api_key=st.session_state['api_key'])

                    ctx = st.session_state['texto_analisis']
                    
                    # Búsqueda Web
                    contexto_web = ""
                    if usar_internet:
                        with st.status("🌐 Buscando...", expanded=True) as s:
                            q = f"{pir} {st.session_state['origen_dato']}"
                            res_web = buscar_en_web(q)
                            contexto_web = f"\nINFO WEB:\n{res_web}\n"
                            s.update(label="✅ Hecho", state="complete", expanded=False)

                    informe_final = f"# INFORME\nFECHA: {datetime.datetime.now().strftime('%d/%m/%Y')}\nFUENTE: {st.session_state['origen_dato']}\nMOTOR: {PROVEEDOR}\n\n"
                    progreso = st.progress(0)

                    # BUCLE PRINCIPAL
                    for i, tec in enumerate(tecnicas_seleccionadas):
                        
                        # Lógica de Prompt
                        instruccion = "Análisis Estratégico Ejecutivo."
                        if "Táctico" in profundidad:
                            qs = DB_CONOCIMIENTO.get(tec, {}).get("preguntas", [])
                            if qs: instruccion = "Responde:\n" + "\n".join([f"- {p}" for p in qs])
                        elif "Operacional" in profundidad:
                            qs = preguntas_manuales.get(tec, [])
                            if qs: instruccion = "Responde SOLO:\n" + "\n".join([f"- {p}" for p in qs])

                        prompt = f"""
                        ACTÚA COMO: Analista de Inteligencia Estratégica y Experto en Relaciones Internacionales. METODOLOGÍA: {tec}. PIR: {pir}
                        DIRECTRICES: Formato académico, BLUF, citar fuentes.
                        {instruccion}
                        CONTEXTO: {ctx[:60000]} {contexto_web}
                        """

                        texto_gen = ""
                        try:
                            if "Google" in PROVEEDOR:
                                model = genai.GenerativeModel("gemini-2.5-flash")
                                res = model.generate_content(prompt)
                                texto_gen = res.text
                            else:
                                client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key_final)
                                completion = client.chat.completions.create(
                                    model="deepseek/deepseek-r1:free",
                                    messages=[{"role": "user", "content": prompt}]
                                )
                                texto_gen = completion.choices[0].message.content
                        except Exception as e:
                            texto_gen = f"Error generando: {e}"

                        firma = f"\n\n> *Análisis generado vía StratIntel Solutions OS ({PROVEEDOR}) | Metodología: {tec}*"
                        informe_final += f"\n\n## 📌 {tec}\n{texto_gen}{firma}\n\n---\n"
                        progreso.progress((i+1)/len(tecnicas_seleccionadas))
                    
                    st.session_state['res'] = informe_final
                    st.session_state['tecnicas_usadas'] = ", ".join(tecnicas_seleccionadas)
                    st.success("✅ Misión Cumplida")
                    st.rerun()

                except Exception as e: st.error(f"Error Fatal: {e}")

# ==========================================================
# 🏁 VISUALIZACIÓN Y DESCARGAS
# ==========================================================
if 'res' in st.session_state and st.session_state['res']:
    st.markdown("---")
    st.markdown(st.session_state['res'])

    # Mapa Visual (Solo si hay Google Key disponible)
    if 'codigo_dot_mapa' not in st.session_state and st.session_state.get('api_key'):
        st.markdown("---")
        with st.spinner("🛰️ Generando Mapa de Relaciones..."):
            grafo, err = generar_esquema_graphviz(st.session_state['res'], st.session_state['api_key'])
            if grafo: 
                st.session_state['codigo_dot_mapa'] = grafo.source
                st.rerun()
            elif err: st.warning(f"No se pudo generar mapa: {err}")

    if 'codigo_dot_mapa' in st.session_state:
        st.subheader("🕸️ Mapa de Relaciones")
        st.graphviz_chart(st.session_state['codigo_dot_mapa'], use_container_width=True)

    # Botones Descarga
    st.markdown("### 📥 Exportar")
    c1, c2 = st.columns(2)
    c1.download_button("Descargar Word", crear_word(st.session_state['res'], st.session_state.get('tecnicas_usadas',''), st.session_state['origen_dato']), "Reporte.docx", use_container_width=True)
    try:
        c2.download_button("Descargar PDF", bytes(crear_pdf(st.session_state['res'], st.session_state.get('tecnicas_usadas',''), st.session_state['origen_dato'])), "Reporte.pdf", use_container_width=True)
    except: pass





